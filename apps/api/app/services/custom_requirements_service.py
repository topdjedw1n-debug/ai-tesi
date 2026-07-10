"""
Custom requirements file upload and processing service
Handles PDF, DOCX, and TXT file uploads for additional document requirements
"""

import io
import logging
from pathlib import Path

from docx import Document as DocxDocument
from fastapi import UploadFile
from pypdf import PdfReader

from app.core.exceptions import ValidationError
from app.services.file_validator import FileValidator

logger = logging.getLogger(__name__)

# Allowed MIME types
ALLOWED_MIME_TYPES = {
    "application/pdf": ".pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    "text/plain": ".txt",
}

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB
MAX_STORED_REQUIREMENTS_CHARS = 50_000
MAX_GENERATION_REQUIREMENTS_CHARS = 60_000
UPLOADED_REQUIREMENTS_START = "=== BEGIN UPLOADED UNIVERSITY REQUIREMENTS ==="
UPLOADED_REQUIREMENTS_END = "=== END UPLOADED UNIVERSITY REQUIREMENTS ==="
RUN_REQUIREMENTS_START = "=== BEGIN GENERATION REQUEST ADDITIONS ==="
RUN_REQUIREMENTS_END = "=== END GENERATION REQUEST ADDITIONS ==="


def combine_generation_requirements(
    persisted_text: str | None, request_text: str | None
) -> str | None:
    """Keep durable intake/methodology context and append per-run additions."""
    persisted = (persisted_text or "").strip()
    requested = (request_text or "").strip()

    if not requested or requested == persisted:
        combined = persisted or None
    elif persisted and requested.startswith(f"{persisted}\n\n"):
        combined = requested
    elif not persisted:
        combined = requested
    else:
        combined = (
            f"{persisted}\n\n"
            f"{RUN_REQUIREMENTS_START}\n"
            f"{requested}\n"
            f"{RUN_REQUIREMENTS_END}"
        )

    if combined and len(combined) > MAX_GENERATION_REQUIREMENTS_CHARS:
        raise ValidationError(
            "Combined methodology and run requirements exceed the "
            f"{MAX_GENERATION_REQUIREMENTS_CHARS}-character generation limit"
        )
    return combined


class CustomRequirementsService:
    """Service for handling custom requirements file uploads"""

    @staticmethod
    def merge_with_existing(existing_text: str | None, uploaded_text: str) -> str:
        """Append uploaded requirements without silently truncating either source."""
        uploaded_text = uploaded_text.strip()
        if not uploaded_text:
            raise ValidationError("Uploaded requirements file contains no usable text")

        uploaded_block = (
            f"{UPLOADED_REQUIREMENTS_START}\n"
            f"{uploaded_text}\n"
            f"{UPLOADED_REQUIREMENTS_END}"
        )
        existing_text = (existing_text or "").strip()
        combined_text = (
            f"{existing_text}\n\n{uploaded_block}" if existing_text else uploaded_block
        )

        if len(combined_text) > MAX_STORED_REQUIREMENTS_CHARS:
            raise ValidationError(
                "Combined intake and uploaded requirements exceed the "
                f"{MAX_STORED_REQUIREMENTS_CHARS}-character limit"
            )

        return combined_text

    @staticmethod
    async def validate_file(file: UploadFile) -> None:
        """
        Validate uploaded file
        Raises ValidationError if file is invalid
        """
        # Check file size
        if file.size and file.size > MAX_FILE_SIZE:
            raise ValidationError(
                f"File size exceeds maximum allowed size of {MAX_FILE_SIZE / 1024 / 1024} MB"
            )

        # Check content type
        if file.content_type not in ALLOWED_MIME_TYPES:
            raise ValidationError(
                f"Unsupported file type: {file.content_type}. "
                f"Allowed types: {', '.join(ALLOWED_MIME_TYPES.keys())}"
            )

        # Verify extension matches content type
        if file.filename:
            ext = Path(file.filename).suffix.lower()
            if ext not in ALLOWED_EXTENSIONS:
                raise ValidationError(
                    f"Unsupported file extension: {ext}. Allowed extensions: {', '.join(ALLOWED_EXTENSIONS)}"
                )
            if ext != ALLOWED_MIME_TYPES[file.content_type]:
                raise ValidationError("File extension does not match content type")

        # Validate file content with magic bytes check
        await FileValidator.validate_file_content(file, file.content_type)

        # Check for ZIP bomb attacks
        await FileValidator.check_zip_bomb(file)

    @staticmethod
    async def extract_text_from_pdf(file: UploadFile) -> str:
        """
        Extract text from PDF file
        Returns extracted text content
        """
        try:
            content = await file.read()
            file.file.seek(0)  # Reset file pointer

            pdf = PdfReader(io.BytesIO(content))
            text_parts = []

            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(page_text)

            extracted_text = "\n\n".join(text_parts)
            if not extracted_text or not extracted_text.strip():
                raise ValidationError(
                    "PDF file appears to be empty or contains no extractable text"
                )

            return extracted_text
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}", exc_info=True)
            raise ValidationError(f"Failed to extract text from PDF: {str(e)}") from e

    @staticmethod
    async def extract_text_from_docx(file: UploadFile) -> str:
        """
        Extract text from DOCX file
        Returns extracted text content
        """
        try:
            content = await file.read()
            file.file.seek(0)  # Reset file pointer

            doc = DocxDocument(io.BytesIO(content))
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # University templates often keep decisive formatting/citation
            # rules in tables, headers, or footers rather than body paragraphs.
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text.strip():
                                text_parts.append(paragraph.text)
            for section in doc.sections:
                for container in (section.header, section.footer):
                    for paragraph in container.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(paragraph.text)

            extracted_text = "\n\n".join(text_parts)
            if not extracted_text or not extracted_text.strip():
                raise ValidationError(
                    "DOCX file appears to be empty or contains no extractable text"
                )

            return extracted_text
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}", exc_info=True)
            raise ValidationError(f"Failed to extract text from DOCX: {str(e)}") from e

    @staticmethod
    async def extract_text_from_txt(file: UploadFile) -> str:
        """
        Extract text from TXT file
        Returns extracted text content
        """
        try:
            content = await file.read()
            file.file.seek(0)  # Reset file pointer

            # Try UTF-8 first, fallback to latin-1
            try:
                text = content.decode("utf-8")
            except UnicodeDecodeError:
                text = content.decode("latin-1")

            if not text or not text.strip():
                raise ValidationError("TXT file appears to be empty")

            return text
        except Exception as e:
            logger.error(f"Error extracting text from TXT: {e}", exc_info=True)
            raise ValidationError(f"Failed to extract text from TXT: {str(e)}") from e

    @staticmethod
    async def extract_text(file: UploadFile) -> str:
        """
        Extract text from uploaded file based on its type
        Returns extracted text content
        """
        # Validate first
        await CustomRequirementsService.validate_file(file)

        # Extract based on file type
        content_type = file.content_type
        if content_type == "application/pdf":
            return await CustomRequirementsService.extract_text_from_pdf(file)
        elif (
            content_type
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            return await CustomRequirementsService.extract_text_from_docx(file)
        elif content_type == "text/plain":
            return await CustomRequirementsService.extract_text_from_txt(file)
        else:
            raise ValidationError(f"Unsupported file type: {content_type}")
