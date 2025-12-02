"""
Document service for managing documents and sections
"""

import io
import logging
from datetime import datetime, timedelta
from typing import Any

from sqlalchemy import delete, func, select, update
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.core.exceptions import NotFoundError, ValidationError
from app.models.auth import User
from app.models.document import Document, DocumentSection
from app.models.payment import Payment

logger = logging.getLogger(__name__)


class DocumentService:
    """Service for document management"""

    def __init__(self, db: AsyncSession):
        self.db = db

    async def check_document_ownership(
        self, document_id: int, user_id: int
    ) -> Document:
        """
        Check if document exists and belongs to user.
        Returns Document object or raises NotFoundError.

        This function ensures IDOR protection by returning 404
        instead of 403 to avoid revealing existence of documents.
        """
        result = await self.db.execute(
            select(Document).where(Document.id == document_id)
        )
        document = result.scalar_one_or_none()

        if not document or document.user_id != user_id:
            raise NotFoundError("Document not found")

        return document

    async def create_document(
        self,
        user_id: int,
        title: str,
        topic: str,
        language: str = "en",
        target_pages: int = 10,
        ai_provider: str = "openai",
        ai_model: str = "gpt-4",
        additional_requirements: str | None = None,
    ) -> dict[str, Any]:
        """Create a new document"""
        try:
            document = Document(
                user_id=user_id,
                title=title,
                topic=topic,
                language=language,
                target_pages=target_pages,
                ai_provider=ai_provider,
                ai_model=ai_model,
                status="draft",
            )

            self.db.add(document)
            await self.db.flush()  # Get the document ID

            # Update user's document count
            await self.db.execute(
                update(User)
                .where(User.id == user_id)
                .values(total_documents_created=User.total_documents_created + 1)
            )

            await self.db.commit()

            # Calculate word count and reading time from document content/title/topic
            content_text = ""
            if document.content:
                content_text = document.content
            elif document.title and document.topic:
                content_text = f"{document.title} {document.topic}"

            word_count = len(content_text.split()) if content_text else 0
            estimated_reading_time = max(
                1, word_count // 200
            )  # Assume 200 WPM reading speed

            return {
                "id": document.id,
                "user_id": document.user_id,
                "title": document.title,
                "topic": document.topic,
                "language": document.language,
                "target_pages": document.target_pages,
                "status": document.status,
                "ai_provider": document.ai_provider,
                "ai_model": document.ai_model,
                "outline": document.outline,
                "content": document.content,
                "tokens_used": document.tokens_used,
                "generation_time_seconds": document.generation_time_seconds,
                "created_at": document.created_at.isoformat(),
                "updated_at": document.updated_at.isoformat()
                if document.updated_at
                else None,
                "completed_at": document.completed_at.isoformat()
                if document.completed_at
                else None,
                "is_archived": document.is_archived,
                "word_count": word_count,
                "estimated_reading_time": estimated_reading_time,
                "sections": [],
            }

        except Exception as e:
            await self.db.rollback()
            logger.error(f"Error creating document: {e}")
            raise ValidationError(f"Failed to create document: {str(e)}") from e

    async def get_document(self, document_id: int, user_id: int) -> dict[str, Any]:
        """Get document by ID"""
        try:
            result = await self.db.execute(
                select(Document)
                .options(selectinload(Document.sections))
                .where(Document.id == document_id, Document.user_id == user_id)
            )
            document = result.scalar_one_or_none()

            if not document:
                raise NotFoundError("Document not found")

            # Calculate word count and reading time from document content/title/topic
            content_text = ""
            if document.content:
                content_text = document.content
            elif document.title and document.topic:
                content_text = f"{document.title} {document.topic}"

            word_count = len(content_text.split()) if content_text else 0
            estimated_reading_time = max(
                1, word_count // 200
            )  # Assume 200 WPM reading speed

            return {
                "id": document.id,
                "user_id": document.user_id,
                "title": document.title,
                "topic": document.topic,
                "language": document.language,
                "target_pages": document.target_pages,
                "status": document.status,
                "is_archived": document.is_archived,
                "ai_provider": document.ai_provider,
                "ai_model": document.ai_model,
                "outline": document.outline,
                "content": document.content,
                "tokens_used": document.tokens_used,
                "generation_time_seconds": document.generation_time_seconds,
                "created_at": document.created_at.isoformat(),
                "updated_at": document.updated_at.isoformat(),
                "completed_at": document.completed_at.isoformat()
                if document.completed_at
                else None,
                "word_count": word_count,
                "estimated_reading_time": estimated_reading_time,
                "sections": [
                    {
                        "id": section.id,
                        "title": section.title,
                        "section_index": section.section_index,
                        "section_type": section.section_type,
                        "content": section.content,
                        "word_count": section.word_count,
                        "status": section.status,
                        "tokens_used": section.tokens_used,
                        "generation_time_seconds": section.generation_time_seconds,
                        "created_at": section.created_at.isoformat(),
                        "completed_at": section.completed_at.isoformat()
                        if section.completed_at
                        else None,
                    }
                    for section in document.sections
                ],
            }

        except Exception as e:
            logger.error(f"Error getting document: {e}")
            raise NotFoundError(f"Failed to get document: {str(e)}") from e

    async def get_user_documents(
        self, user_id: int, limit: int = 20, offset: int = 0
    ) -> dict[str, Any]:
        """Get user's documents with pagination"""
        try:
            # Get total count
            count_result = await self.db.execute(
                select(Document.id).where(Document.user_id == user_id)
            )
            total_count = len(count_result.scalars().all())

            # Get documents
            result = await self.db.execute(
                select(Document)
                .where(Document.user_id == user_id)
                .order_by(Document.created_at.desc())
                .limit(limit)
                .offset(offset)
            )
            documents = result.scalars().all()

            # Calculate pagination metadata
            total_pages = (total_count + limit - 1) // limit if limit > 0 else 1
            page = (offset // limit) + 1 if limit > 0 else 1

            # Build document list with all required fields
            doc_list = []
            for doc in documents:
                # Calculate word count and reading time
                content_text = ""
                if doc.content:
                    content_text = doc.content
                elif doc.title and doc.topic:
                    content_text = f"{doc.title} {doc.topic}"

                word_count = len(content_text.split()) if content_text else 0
                estimated_reading_time = max(1, word_count // 200)

                doc_list.append(
                    {
                        "id": doc.id,
                        "user_id": doc.user_id,
                        "title": doc.title,
                        "topic": doc.topic,
                        "language": doc.language,
                        "target_pages": doc.target_pages,
                        "status": doc.status,
                        "is_archived": doc.is_archived,
                        "ai_provider": doc.ai_provider,
                        "ai_model": doc.ai_model,
                        "outline": doc.outline,
                        "content": doc.content,
                        "tokens_used": doc.tokens_used,
                        "generation_time_seconds": doc.generation_time_seconds,
                        "created_at": doc.created_at.isoformat(),
                        "updated_at": doc.updated_at.isoformat()
                        if doc.updated_at
                        else None,
                        "word_count": word_count,
                        "estimated_reading_time": estimated_reading_time,
                        "sections": [],
                    }
                )

            return {
                "documents": doc_list,
                "total": total_count,
                "page": page,
                "per_page": limit,
                "total_pages": total_pages,
            }

        except Exception as e:
            logger.error(f"Error getting user documents: {e}")
            raise ValidationError(f"Failed to get documents: {str(e)}") from e

    async def get_user_stats(self, user_id: int) -> dict[str, Any]:
        """
        Get user statistics for dashboard.

        Returns:
            dict with:
                - totalDocuments: int
                - totalWords: int
                - totalCost: float (from payments)
                - totalTokens: int
        """
        try:
            # Get total documents count
            doc_count_result = await self.db.execute(
                select(func.count(Document.id)).where(Document.user_id == user_id)
            )
            total_documents = doc_count_result.scalar() or 0

            # Get total words from all documents
            # Calculate word count from content
            docs_result = await self.db.execute(
                select(Document).where(Document.user_id == user_id)
            )
            documents = docs_result.scalars().all()

            total_words = 0
            total_tokens = 0
            for doc in documents:
                # Calculate word count
                content_text = ""
                if doc.content:
                    content_text = str(doc.content)
                elif doc.title and doc.topic:
                    content_text = f"{doc.title} {doc.topic}"

                word_count = len(content_text.split()) if content_text else 0
                total_words += word_count

                # Sum tokens
                total_tokens += int(doc.tokens_used or 0)

            # Get total cost from payments
            payments_result = await self.db.execute(
                select(func.sum(Payment.amount)).where(
                    Payment.user_id == user_id, Payment.status == "completed"
                )
            )
            total_cost = float(payments_result.scalar() or 0)

            return {
                "totalDocuments": total_documents,
                "totalWords": total_words,
                "totalCost": total_cost,
                "totalTokens": total_tokens,
            }
        except Exception as e:
            logger.error(f"Error getting user stats: {e}")
            raise ValidationError(f"Failed to get user stats: {str(e)}") from e

    async def get_recent_activity(
        self, user_id: int, limit: int = 10
    ) -> list[dict[str, Any]]:
        """
        Get recent activity for user dashboard.

        Returns list of activities based on document status changes.
        """
        try:
            # Get recent documents ordered by updated_at
            result = await self.db.execute(
                select(Document)
                .where(Document.user_id == user_id)
                .order_by(Document.updated_at.desc())
                .limit(limit)
            )
            documents = result.scalars().all()

            activities = []
            for doc in documents:
                # Map document status to activity type
                activity_type = None
                activity_status = "success"

                if doc.status == "completed":
                    activity_type = "document_completed"
                elif doc.status == "sections_generated":
                    activity_type = "section_generated"
                elif doc.status == "outline_generated":
                    activity_type = "outline_generated"
                elif doc.status == "draft":
                    activity_type = "document_created"
                elif doc.status == "generating":
                    activity_type = "section_generated"
                    activity_status = "pending"
                elif doc.status == "failed":
                    activity_type = "section_generated"
                    activity_status = "error"
                elif doc.status == "payment_pending":
                    activity_type = "document_created"
                    activity_status = "pending"

                if activity_type:
                    # Get description based on status
                    description = ""
                    if doc.status == "completed":
                        word_count = len(doc.content.split()) if doc.content else 0
                        description = f"Thesis completed with {word_count:,} words"
                    elif doc.status == "sections_generated":
                        description = "Sections generated successfully"
                    elif doc.status == "outline_generated":
                        description = "Complete outline generated"
                    elif doc.status == "draft":
                        description = "Document created"
                    elif doc.status == "generating":
                        description = "Generation in progress"
                    elif doc.status == "failed":
                        description = "Generation failed"
                    elif doc.status == "payment_pending":
                        description = "Payment pending"

                    activities.append(
                        {
                            "id": doc.id,
                            "type": activity_type,
                            "title": doc.title or f"Document {doc.id}",
                            "description": description,
                            "timestamp": doc.updated_at.isoformat()
                            if doc.updated_at
                            else doc.created_at.isoformat(),
                            "status": activity_status,
                            "document_id": doc.id,
                        }
                    )

            return activities
        except Exception as e:
            logger.error(f"Error getting recent activity: {e}")
            raise ValidationError(f"Failed to get recent activity: {str(e)}") from e

    async def update_document(
        self, document_id: int, user_id: int, **updates: Any
    ) -> dict[str, Any]:
        """Update document"""
        try:
            # Check if document exists and belongs to user
            result = await self.db.execute(
                select(Document).where(
                    Document.id == document_id, Document.user_id == user_id
                )
            )
            document = result.scalar_one_or_none()

            if not document:
                raise NotFoundError("Document not found")

            # Update allowed fields
            allowed_fields = [
                "title",
                "topic",
                "language",
                "target_pages",
                "ai_provider",
                "ai_model",
                "is_public",
            ]

            update_data = {k: v for k, v in updates.items() if k in allowed_fields}

            if update_data:
                await self.db.execute(
                    update(Document)
                    .where(Document.id == document_id)
                    .values(**update_data)
                )
                await self.db.commit()

            return {"message": "Document updated successfully"}

        except NotFoundError:
            await self.db.rollback()
            raise
        except Exception as e:
            await self.db.rollback()
            logger.error(f"Error updating document: {e}")
            raise ValidationError(f"Failed to update document: {str(e)}") from e

    async def delete_document(self, document_id: int, user_id: int) -> dict[str, Any]:
        """Delete document"""
        try:
            # Check if document exists and belongs to user
            result = await self.db.execute(
                select(Document).where(
                    Document.id == document_id, Document.user_id == user_id
                )
            )
            document = result.scalar_one_or_none()

            if not document:
                raise NotFoundError("Document not found")

            # Delete document (cascade will handle sections)
            await self.db.execute(delete(Document).where(Document.id == document_id))

            await self.db.commit()

            return {"message": "Document deleted successfully"}

        except Exception as e:
            await self.db.rollback()
            logger.error(f"Error deleting document: {e}")
            raise ValidationError(f"Failed to delete document: {str(e)}") from e

    async def get_document_sections(
        self, document_id: int, user_id: int
    ) -> list[dict[str, Any]]:
        """Get all sections for a document"""
        try:
            # Verify document ownership
            result = await self.db.execute(
                select(Document.id).where(
                    Document.id == document_id, Document.user_id == user_id
                )
            )
            if not result.scalar_one_or_none():
                raise NotFoundError("Document not found")

            # Get sections
            result = await self.db.execute(
                select(DocumentSection)
                .where(DocumentSection.document_id == document_id)
                .order_by(DocumentSection.section_index)
            )
            sections = result.scalars().all()

            return [
                {
                    "id": section.id,
                    "title": section.title,
                    "section_index": section.section_index,
                    "section_type": section.section_type,
                    "content": section.content,
                    "word_count": section.word_count,
                    "status": section.status,
                    "tokens_used": section.tokens_used,
                    "generation_time_seconds": section.generation_time_seconds,
                    "created_at": section.created_at.isoformat(),
                    "completed_at": section.completed_at.isoformat()
                    if section.completed_at
                    else None,
                }
                for section in sections
            ]

        except Exception as e:
            logger.error(f"Error getting document sections: {e}")
            raise NotFoundError(f"Failed to get document sections: {str(e)}") from e

    async def update_document_content(
        self, document_id: int, user_id: int, content: str
    ) -> dict[str, Any]:
        """Update document content"""
        try:
            # Verify document ownership
            result = await self.db.execute(
                select(Document).where(
                    Document.id == document_id, Document.user_id == user_id
                )
            )
            document = result.scalar_one_or_none()

            if not document:
                raise NotFoundError("Document not found")

            # Update content
            await self.db.execute(
                update(Document)
                .where(Document.id == document_id)
                .values(content=content, status="completed")
            )

            await self.db.commit()

            return {"message": "Document content updated successfully"}

        except Exception as e:
            await self.db.rollback()
            logger.error(f"Error updating document content: {e}")
            raise ValidationError(f"Failed to update document content: {str(e)}") from e

    async def verify_file_storage_integrity(self) -> dict[str, Any]:
        """
        Verify file storage integrity (MinIO/S3).

        Checks:
        - Files referenced in database exist in storage
        - Orphaned files in storage (not referenced in DB)
        - File accessibility

        Returns:
            dict with missing_files count, orphaned_files list, and status
        """
        import time

        from minio import Minio
        from minio.error import S3Error

        from app.core.config import settings
        from app.services.storage_service import StorageService

        missing_files = []
        orphaned_files = []
        storage_files = []

        try:
            # Initialize MinIO client for list_objects (StorageService doesn't have this yet)
            client = Minio(
                settings.MINIO_ENDPOINT,
                access_key=settings.MINIO_ACCESS_KEY,
                secret_key=settings.MINIO_SECRET_KEY,
                secure=settings.MINIO_SECURE,
            )
            
            # Initialize StorageService for file_exists checks
            storage_service = StorageService()

            # Get all documents with file paths
            result = await self.db.execute(
                select(Document).where(
                    (Document.docx_path.isnot(None)) | (Document.pdf_path.isnot(None))
                )
            )
            documents = result.scalars().all()

            # Check files referenced in DB exist in storage using StorageService
            for doc in documents:
                if doc.docx_path:
                    exists = await storage_service.file_exists(doc.docx_path)
                    if not exists:
                        missing_files.append(
                            {"document_id": doc.id, "file": doc.docx_path}
                        )

                if doc.pdf_path:
                    exists = await storage_service.file_exists(doc.pdf_path)
                    if not exists:
                        missing_files.append(
                            {"document_id": doc.id, "file": doc.pdf_path}
                        )

            # List all files in bucket
            try:
                objects = client.list_objects(settings.MINIO_BUCKET, recursive=True)
                for obj in objects:
                    storage_files.append(obj.object_name)

                    # Check if file is referenced in DB
                    is_referenced = False
                    for doc in documents:
                        if doc.docx_path and obj.object_name in doc.docx_path:
                            is_referenced = True
                            break
                        if doc.pdf_path and obj.object_name in doc.pdf_path:
                            is_referenced = True
                            break

                    if not is_referenced:
                        orphaned_files.append(obj.object_name)
            except S3Error as e:
                logger.error(f"Error listing storage objects: {e}")

            # Determine status
            if len(missing_files) == 0 and len(orphaned_files) == 0:
                status = "healthy"
            elif len(missing_files) > 0:
                status = "critical"
            else:
                status = "needs_attention"

            return {
                "status": status,
                "missing_files": len(missing_files),
                "missing_file_details": missing_files[:10],  # Limit details
                "orphaned_files": len(orphaned_files),
                "orphaned_file_details": orphaned_files[:10],  # Limit details
                "total_storage_files": len(storage_files),
                "timestamp": time.time(),
            }
        except Exception as e:
            logger.error(f"Storage verification failed: {e}")
            return {
                "status": "error",
                "missing_files": 0,
                "orphaned_files": 0,
                "error": str(e),
                "timestamp": time.time(),
            }

    async def export_document(
        self, document_id: int, format: str, user_id: int
    ) -> dict[str, Any]:
        """Export document to DOCX or PDF format"""
        logger.info(f"🚀 Starting export_document: doc_id={document_id}, format={format}, user_id={user_id}")
        try:
            from docx import Document as DocxDocument
            from minio import Minio
            from minio.error import S3Error

            from app.core.config import settings

            logger.info(f"📄 Getting document {document_id} from database...")
            # Get document with sections
            result = await self.db.execute(
                select(Document)
                .options(selectinload(Document.sections))
                .where(Document.id == document_id, Document.user_id == user_id)
            )
            document = result.scalar_one_or_none()

            if not document:
                logger.error(f"❌ Document {document_id} not found for user {user_id}")
                raise NotFoundError("Document not found")

            # Check if document has content or sections
            if document.status not in ["completed", "sections_generated"]:
                raise ValidationError(
                    "Document is not ready for export. Status must be 'completed' or 'sections_generated'."
                )

            # Generate file based on format
            if format == "docx":
                # Create DOCX document
                docx = DocxDocument()

                # Add title
                docx.add_heading(document.title, 0)

                # Add metadata
                docx.add_paragraph(f"Topic: {document.topic}")
                docx.add_paragraph(f"Language: {document.language}")
                docx.add_paragraph(
                    f"Created: {document.created_at.strftime('%Y-%m-%d %H:%M:%S')}"
                )
                docx.add_paragraph("")  # Empty line

                # Add content if available
                if document.content:
                    docx.add_paragraph(document.content)
                elif document.sections:
                    # Add sections
                    for section in sorted(
                        document.sections, key=lambda s: s.section_index
                    ):
                        docx.add_heading(section.title, 1)
                        if section.content:
                            docx.add_paragraph(section.content)
                        docx.add_paragraph("")  # Empty line between sections

                # Save to BytesIO
                file_stream = io.BytesIO()
                logger.info(f"Saving DOCX document for doc_id={document_id}")
                docx.save(file_stream)
                file_size = file_stream.tell()  # Get size BEFORE seeking to 0
                logger.info(f"DOCX saved, file_size after save: {file_size} bytes")
                file_stream.seek(0)
                file_data = file_stream.getvalue()  # Get bytes for upload

                file_extension = "docx"
                content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

            elif format == "pdf":
                # Create PDF document using ReportLab
                from reportlab.lib import colors
                from reportlab.lib.pagesizes import letter
                from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
                from reportlab.lib.units import inch
                from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

                # Create PDF in memory
                file_stream = io.BytesIO()
                pdf = SimpleDocTemplate(
                    file_stream,
                    pagesize=letter,
                    rightMargin=72,
                    leftMargin=72,
                    topMargin=72,
                    bottomMargin=18,
                )

                # Container for PDF elements
                elements = []
                styles = getSampleStyleSheet()

                # Custom styles
                title_style = ParagraphStyle(
                    "CustomTitle",
                    parent=styles["Heading1"],
                    fontSize=24,
                    textColor=colors.HexColor("#1a1a1a"),
                    spaceAfter=30,
                    alignment=1,  # Center
                )

                heading_style = ParagraphStyle(
                    "CustomHeading",
                    parent=styles["Heading2"],
                    fontSize=16,
                    textColor=colors.HexColor("#2c3e50"),
                    spaceAfter=12,
                )

                body_style = ParagraphStyle(
                    "CustomBody",
                    parent=styles["Normal"],
                    fontSize=11,
                    leading=14,
                    spaceAfter=12,
                )

                # Add title
                elements.append(Paragraph(document.title, title_style))
                elements.append(Spacer(1, 0.2 * inch))

                # Add metadata
                metadata = f"""
                <b>Topic:</b> {document.topic}<br/>
                <b>Language:</b> {document.language}<br/>
                <b>Created:</b> {document.created_at.strftime('%Y-%m-%d %H:%M:%S')}
                """
                elements.append(Paragraph(metadata, body_style))
                elements.append(Spacer(1, 0.3 * inch))

                # Add content
                if document.content:
                    # Split content by paragraphs
                    for paragraph in document.content.split("\n\n"):
                        if paragraph.strip():
                            # Escape HTML special characters
                            safe_text = (
                                paragraph.replace("&", "&amp;")
                                .replace("<", "&lt;")
                                .replace(">", "&gt;")
                            )
                            elements.append(Paragraph(safe_text, body_style))
                elif document.sections:
                    # Add sections
                    for section in sorted(
                        document.sections, key=lambda s: s.section_index
                    ):
                        # Section heading
                        elements.append(Paragraph(section.title, heading_style))
                        elements.append(Spacer(1, 0.1 * inch))

                        # Section content
                        if section.content:
                            for paragraph in section.content.split("\n\n"):
                                if paragraph.strip():
                                    safe_text = (
                                        paragraph.replace("&", "&amp;")
                                        .replace("<", "&lt;")
                                        .replace(">", "&gt;")
                                    )
                                    elements.append(Paragraph(safe_text, body_style))

                        elements.append(Spacer(1, 0.2 * inch))

                # Build PDF
                logger.info(f"Generating PDF document for doc_id={document_id}")
                pdf.build(elements)
                file_size = file_stream.tell()
                logger.info(f"PDF generated, file_size: {file_size} bytes")
                file_stream.seek(0)
                file_data = file_stream.getvalue()  # Get bytes for upload

                file_extension = "pdf"
                content_type = "application/pdf"

            else:
                raise ValidationError(f"Unsupported export format: {format}")

            # Generate object name
            object_name = (
                f"documents/{user_id}/{document_id}/{document_id}.{file_extension}"
            )

            logger.info(f"Uploading {format} file: {object_name} ({file_size} bytes)")

            # Upload file using StorageService
            from app.services.storage_service import StorageService
            
            storage_service = StorageService()
            storage_path = await storage_service.upload_file(
                object_name, 
                file_data, 
                content_type
            )
            
            logger.info(f"Successfully uploaded to MinIO: {storage_path}")

            # Generate public download URL (since bucket is public)
            download_url = f"http://{settings.MINIO_ENDPOINT}/{settings.MINIO_BUCKET}/{object_name}"
            logger.info(f"Generated download URL: {download_url}")

            # Update document path
            storage_path = f"s3://{settings.MINIO_BUCKET}/{object_name}"
            if format == "docx":
                await self.db.execute(
                    update(Document)
                    .where(Document.id == document_id)
                    .values(docx_path=storage_path)
                )
            elif format == "pdf":
                await self.db.execute(
                    update(Document)
                    .where(Document.id == document_id)
                    .values(pdf_path=storage_path)
                )

            await self.db.commit()

            # Return response
            return {
                "download_url": download_url,
                "expires_at": (datetime.utcnow() + timedelta(hours=1)).isoformat(),
                "file_size": file_size,
                "format": format,
            }

        except NotFoundError:
            raise
        except ValidationError:
            raise
        except Exception as e:
            await self.db.rollback()
            logger.error(f"Error exporting document: {e}")
            raise ValidationError(f"Failed to export document: {str(e)}") from e
