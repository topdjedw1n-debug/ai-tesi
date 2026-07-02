"""
Stage B2: bibliography must survive the background generation path and reach
the exported DOCX/PDF.

- SectionGenerator builds per-section bibliographies; the background job must
  persist them on DocumentSection (checkpoint-resume loses in-memory results).
- Document assembly appends a merged, deduped "# Bibliografia" section to
  Document.content (language-aware heading).
- Exporters render "# " blocks as real headings, so the bibliography shows up
  in DOCX and PDF.
"""

import io
from contextlib import ExitStack
from unittest.mock import AsyncMock, MagicMock, patch

import pytest
from sqlalchemy import select

import tests.test_provenance_ledger as harness_mod
from app.models.document import Document, DocumentSection
from app.services.ai_pipeline.citation_formatter import (
    bibliography_heading,
    merge_bibliographies,
)
from app.services.background_jobs import BackgroundJobService
from app.services.document_service import DocumentService


@pytest.fixture
def mock_redis():
    redis = MagicMock()
    redis.get = AsyncMock(return_value=None)
    redis.set = AsyncMock()
    redis.delete = AsyncMock()
    return redis


REF_A = (
    "Vaswani, A. (2017). Attention is all you need. NeurIPS. "
    "https://doi.org/10.5555/3295222"
)
REF_B = (
    "Devlin, J. (2019). BERT: Pre-training of deep bidirectional "
    "transformers. NAACL. https://doi.org/10.18653/v1/N19-1423"
)


def section_result_with_bibliography(index, bibliography, pack_keys=None):
    return {
        "section_title": f"Section {index}",
        "section_index": index,
        "content": f"Generated content for section {index}",
        "citations": [],
        "bibliography": bibliography,
        "pack_keys_used": pack_keys or [],
        "sources_used": len(bibliography),
        "humanized": False,
        "cited_sources": [],
    }


# ---------------------------------------------------------------------------
# Unit: merge helper
# ---------------------------------------------------------------------------


def test_merge_bibliographies_dedupes_and_sorts():
    merged = merge_bibliographies(
        [
            [REF_B, REF_A],
            [REF_A, "  "],  # duplicate + blank
            None,  # legacy section without the column
            [],
        ]
    )
    assert merged == sorted([REF_A, REF_B], key=str.casefold)
    assert len(merged) == 2


def test_bibliography_heading_is_language_aware():
    assert bibliography_heading("it") == "Bibliografia"
    assert bibliography_heading("it-IT") == "Bibliografia"
    assert bibliography_heading("en") == "Bibliography"
    assert bibliography_heading("uk") == "Бібліографія"
    assert bibliography_heading(None) == "Bibliography"
    assert bibliography_heading("de") == "Bibliography"  # fallback


# ---------------------------------------------------------------------------
# Pipeline: persistence + document assembly
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_section_save_persists_bibliography_and_pack_keys(
    db_session, mock_redis, monkeypatch
):
    monkeypatch.setattr(
        "app.services.background_jobs.settings", harness_mod.make_settings()
    )
    user, document = await harness_mod.seed_document(
        db_session, section_titles=("Intro", "Methods")
    )
    document_id = document.id

    with ExitStack() as stack:
        harness_mod.pipeline_harness(
            stack,
            db_session,
            mock_redis,
            generate_side_effect=[
                section_result_with_bibliography(1, [REF_A], ["Vaswani2017"]),
                section_result_with_bibliography(2, [REF_B, REF_A], ["Devlin2019"]),
            ],
        )
        await BackgroundJobService.generate_full_document(
            document_id=document_id, user_id=user.id
        )

    sections = (
        (
            await db_session.execute(
                select(DocumentSection)
                .where(DocumentSection.document_id == document_id)
                .order_by(DocumentSection.section_index)
            )
        )
        .scalars()
        .all()
    )
    assert sections[0].bibliography == [REF_A]
    assert sections[0].pack_keys_used == ["Vaswani2017"]
    assert sections[1].bibliography == [REF_B, REF_A]
    assert sections[1].pack_keys_used == ["Devlin2019"]


@pytest.mark.asyncio
@pytest.mark.parametrize(
    ("language", "heading"),
    [("it", "Bibliografia"), ("en", "Bibliography"), ("uk", "Бібліографія")],
)
async def test_document_content_gets_bibliography_section(
    db_session, mock_redis, monkeypatch, language, heading
):
    monkeypatch.setattr(
        "app.services.background_jobs.settings", harness_mod.make_settings()
    )
    user, document = await harness_mod.seed_document(
        db_session, section_titles=("Intro", "Methods")
    )
    document.language = language
    await db_session.commit()
    document_id = document.id

    with ExitStack() as stack:
        harness_mod.pipeline_harness(
            stack,
            db_session,
            mock_redis,
            generate_side_effect=[
                section_result_with_bibliography(1, [REF_B]),
                section_result_with_bibliography(2, [REF_A, REF_B]),  # dup REF_B
            ],
        )
        await BackgroundJobService.generate_full_document(
            document_id=document_id, user_id=user.id
        )

    content = (
        await db_session.execute(
            select(Document.content).where(Document.id == document_id)
        )
    ).scalar_one()
    marker = f"# {heading}"
    assert marker in content
    bibliography_block = content.split(marker, 1)[1]
    # Deduped: each reference exactly once, alphabetical order
    assert bibliography_block.count(REF_A) == 1
    assert bibliography_block.count(REF_B) == 1
    assert bibliography_block.index(REF_B) < bibliography_block.index(REF_A)
    # DOI strings from the pack land in the final content
    assert "https://doi.org/" in bibliography_block


@pytest.mark.asyncio
async def test_empty_bibliography_appends_nothing(db_session, mock_redis, monkeypatch):
    monkeypatch.setattr(
        "app.services.background_jobs.settings", harness_mod.make_settings()
    )
    user, document = await harness_mod.seed_document(
        db_session, section_titles=("Only",)
    )
    document_id = document.id

    with ExitStack() as stack:
        harness_mod.pipeline_harness(
            stack,
            db_session,
            mock_redis,
            generate_side_effect=[section_result_with_bibliography(1, [])],
        )
        await BackgroundJobService.generate_full_document(
            document_id=document_id, user_id=user.id
        )

    content = (
        await db_session.execute(
            select(Document.content).where(Document.id == document_id)
        )
    ).scalar_one()
    assert "Bibliograf" not in content
    assert "Бібліографія" not in content


@pytest.mark.asyncio
async def test_regen_persists_final_attempt_bibliography(
    db_session, mock_redis, monkeypatch
):
    """Attempt 1 fails a gate with bibliography X; accepted attempt 2 has Y —
    only Y may be persisted and exported."""
    monkeypatch.setattr(
        "app.services.background_jobs.settings",
        harness_mod.make_settings(
            QUALITY_GATES_ENABLED=True, QUALITY_MAX_REGENERATE_ATTEMPTS=1
        ),
    )
    user, document = await harness_mod.seed_document(
        db_session, section_titles=("Only",)
    )
    document_id = document.id

    with ExitStack() as stack:
        harness_mod.pipeline_harness(
            stack,
            db_session,
            mock_redis,
            generate_side_effect=[
                section_result_with_bibliography(1, [REF_A], ["Vaswani2017"]),
                section_result_with_bibliography(1, [REF_B], ["Devlin2019"]),
            ],
            grammar_side_effect=[
                (40.0, 25, "failed", "Too many grammar errors"),
                (95.0, 0, "passed", None),
            ],
        )
        await BackgroundJobService.generate_full_document(
            document_id=document_id, user_id=user.id
        )

    section = (
        await db_session.execute(
            select(DocumentSection).where(DocumentSection.document_id == document_id)
        )
    ).scalar_one()
    assert section.bibliography == [REF_B]
    assert section.pack_keys_used == ["Devlin2019"]
    content = (
        await db_session.execute(
            select(Document.content).where(Document.id == document_id)
        )
    ).scalar_one()
    assert REF_B in content
    assert REF_A not in content


# ---------------------------------------------------------------------------
# Exporters
# ---------------------------------------------------------------------------


async def _seed_completed_document(db_session, *, content=None, sections=()):
    from app.models.auth import User

    user = User(email="biblio-export@example.com", full_name="Biblio Exporter")
    db_session.add(user)
    await db_session.commit()
    await db_session.refresh(user)

    document = Document(
        user_id=user.id,
        title="Export Test Thesis",
        topic="AI in Education",
        status="completed",
        language="it",
        content=content,
    )
    db_session.add(document)
    await db_session.commit()
    await db_session.refresh(document)

    for index, (title, text, bibliography) in enumerate(sections, start=1):
        db_session.add(
            DocumentSection(
                document_id=document.id,
                title=title,
                section_index=index,
                content=text,
                status="completed",
                bibliography=bibliography,
            )
        )
    await db_session.commit()
    return user, document


def _capture_upload(stack):
    """Patch StorageService.upload_file and capture the uploaded bytes."""
    uploads = {}

    async def fake_upload(object_name, file_data, content_type):
        uploads["object_name"] = object_name
        uploads["data"] = file_data
        uploads["content_type"] = content_type
        return f"s3://test/{object_name}"

    storage_class = stack.enter_context(
        patch("app.services.storage_service.StorageService")
    )
    storage_class.return_value.upload_file = AsyncMock(side_effect=fake_upload)
    return uploads


@pytest.mark.asyncio
async def test_export_docx_renders_bibliografia_heading(db_session):
    content = (
        "# Introduzione\n\nTesto della sezione.\n\n"
        f"# Bibliografia\n\n{REF_A}\n\n{REF_B}"
    )
    user, document = await _seed_completed_document(db_session, content=content)

    with ExitStack() as stack:
        uploads = _capture_upload(stack)
        result = await DocumentService(db_session).export_document(
            document_id=document.id, format="docx", user_id=user.id
        )

    assert result["format"] == "docx"
    from docx import Document as DocxDocument

    docx = DocxDocument(io.BytesIO(uploads["data"]))
    headings = [p.text for p in docx.paragraphs if p.style.name.startswith("Heading")]
    assert "Introduzione" in headings
    assert "Bibliografia" in headings
    body_text = "\n".join(p.text for p in docx.paragraphs)
    assert "https://doi.org/10.5555/3295222" in body_text
    assert "https://doi.org/10.18653/v1/N19-1423" in body_text


@pytest.mark.asyncio
async def test_export_docx_sections_fallback_appends_bibliography(db_session):
    """No document.content: bibliography is rebuilt from persisted
    per-section references."""
    user, document = await _seed_completed_document(
        db_session,
        content=None,
        sections=(
            ("Introduzione", "Testo uno.", [REF_A]),
            ("Metodi", "Testo due.", [REF_B, REF_A]),
        ),
    )

    with ExitStack() as stack:
        uploads = _capture_upload(stack)
        await DocumentService(db_session).export_document(
            document_id=document.id, format="docx", user_id=user.id
        )

    from docx import Document as DocxDocument

    docx = DocxDocument(io.BytesIO(uploads["data"]))
    headings = [p.text for p in docx.paragraphs if p.style.name.startswith("Heading")]
    assert "Bibliografia" in headings
    body_text = "\n".join(p.text for p in docx.paragraphs)
    # Deduped: REF_A appears once despite being cited by both sections
    assert body_text.count("https://doi.org/10.5555/3295222") == 1


@pytest.mark.asyncio
async def test_export_docx_empty_bibliography_no_heading_no_crash(db_session):
    user, document = await _seed_completed_document(
        db_session,
        content=None,
        sections=(("Introduzione", "Testo.", []), ("Metodi", "Altro testo.", None)),
    )

    with ExitStack() as stack:
        uploads = _capture_upload(stack)
        result = await DocumentService(db_session).export_document(
            document_id=document.id, format="docx", user_id=user.id
        )

    assert result["file_size"] > 0
    from docx import Document as DocxDocument

    docx = DocxDocument(io.BytesIO(uploads["data"]))
    assert all("Bibliografia" not in p.text for p in docx.paragraphs)


@pytest.mark.asyncio
@pytest.mark.parametrize("with_bibliography", [True, False])
async def test_export_pdf_with_bibliography_does_not_crash(
    db_session, with_bibliography
):
    content = "# Introduzione\n\nTesto della sezione."
    if with_bibliography:
        content += f"\n\n# Bibliografia\n\n{REF_A}"
    user, document = await _seed_completed_document(db_session, content=content)

    with ExitStack() as stack:
        _capture_upload(stack)
        result = await DocumentService(db_session).export_document(
            document_id=document.id, format="pdf", user_id=user.id
        )

    assert result["format"] == "pdf"
    assert result["file_size"] > 0
